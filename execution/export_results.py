# PURPOSE:
# This script takes the final JSON data that the AI extracted and turns it 
# into a clean Markdown report that is easy for humans to read.
# It also generates DOCX and PDF files if requested by the configuration.

import json
import os
import sys

# Optional libraries for PDF/DOCX generation
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


def create_markdown(data, topic, source_url):
    report_path = 'outputs/report.md'
    try:
        with open(report_path, 'w', encoding='utf-8') as f:
            # Title & Topic
            f.write("# Executive Research Report\n\n")
            f.write(f"**Research Topic:** {topic}\n")
            f.write(f"**Source Used:** {source_url}\n\n")
            
            # Summary Section
            f.write("## 1. Executive Summary\n")
            f.write("This report contains automatically extracted facts and insights related to your research topic. ")
            f.write("The facts below were pulled directly from the source URL without human intervention.\n\n")
            
            # Findings Section
            f.write("## 2. Key Findings & Topics\n")
            if not data:
                f.write("*No exact facts were found for this topic in the source text.*\n\n")
            else:
                for item in data:
                    fact_text = item.get('fact', 'Unknown fact')
                    f.write(f"- {fact_text}\n")
            
            f.write("\n")
            # Limitations Section
            f.write("## 3. Notes & Limitations\n")
            f.write("- **MVP Notice:** This is a Version 1 AI Research Agent.\n")
            f.write("- **AI Context:** It extracts sentences matching keywords but does not yet summarize them natively.\n")
                    
        return report_path
    except Exception as e:
        print(f"Error creating Markdown: {e}")
        return None


def create_docx(data, topic, source_url):
    if not DOCX_AVAILABLE:
        print("python-docx not installed. Skipping DOCX generation.")
        return None
        
    try:
        doc = Document()
        doc.add_heading('Executive Research Report', 0)
        
        # Topic
        p = doc.add_paragraph()
        run = p.add_run('Research Topic: ')
        run.bold = True
        p.add_run(str(topic))
        
        # Source
        p = doc.add_paragraph()
        run = p.add_run('Source Used: ')
        run.bold = True
        p.add_run(str(source_url))
        
        # Summary
        doc.add_heading('1. Executive Summary', level=1)
        doc.add_paragraph(
            "This report contains automatically extracted facts and "
            "insights related to your research topic."
        )
        
        # Findings
        doc.add_heading('2. Key Findings & Topics', level=1)
        if not data:
             doc.add_paragraph("No exact facts were found for this topic.")
        else:
            for item in data:
                # Using a dash instead of 'List Bullet' style due to template issues
                doc.add_paragraph(f"- {str(item.get('fact', ''))}")
                
        doc.add_heading('3. Notes & Limitations', level=1)
        doc.add_paragraph("- MVP Notice: This is a Version 1 AI Research Agent.")
        
        out_path = 'outputs/report.docx'
        doc.save(out_path)
        return out_path
    except Exception as e:
        print(f"Error generating DOCX: {e}")
        return None


def create_pdf(data, topic, source_url):
    if not PDF_AVAILABLE:
        print("fpdf2 not installed. Skipping PDF generation.")
        return None

    try:
        class CustomPDF(FPDF):
            def header(self):
                self.set_font('helvetica', 'B', 15)
                self.cell(0, 10, 'Executive Research Report', border=False, ln=1, align='C')
                self.ln(5)

        pdf = CustomPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("helvetica", size=12)
        
        # Content Width (Page width minus margins)
        w = pdf.w - 2 * pdf.l_margin
        
        # Topic and Source
        pdf.set_font(style="B")
        pdf.multi_cell(w, 8, f"Topic: {str(topic)}", ln=1)
        pdf.multi_cell(w, 8, f"Source: {str(source_url)}", ln=1)
        pdf.ln(5)
        
        # Summary
        pdf.set_font(style="B", size=14)
        pdf.cell(w, 8, "1. Executive Summary", ln=1)
        pdf.set_font(style="", size=11)
        pdf.multi_cell(w, 6, "This report contains automatically extracted facts and insights related to your research topic.", ln=1)
        pdf.ln(5)
        
        # Findings
        pdf.set_font(style="B", size=14)
        pdf.cell(w, 8, "2. Key Findings & Topics", ln=1)
        pdf.set_font(style="", size=11)
        if not data:
            pdf.cell(w, 8, "No exact facts were found.", ln=1)
        else:
            for item in data:
                # Scrub non-latin characters for the basic PDF font
                fact_str = str(item.get('fact', '')).replace('\n', ' ')
                fact_str = fact_str.encode('ascii', 'ignore').decode('ascii')
                pdf.multi_cell(w, 7, f"- {fact_str}", ln=1)

        # Save
        out_path = 'outputs/report.pdf'
        pdf.output(out_path)
        return out_path
    except Exception as e:
        print(f"Error generating PDF: {e}")
        return None


def export_reports(topic, source_url):
    try:
        # Resolve extracted facts path
        facts_path = os.path.join('.tmp', 'extracted_facts.json')
        with open(facts_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error loading data: {e}")
        return

    # Ensure output directory exists
    os.makedirs('outputs', exist_ok=True)
    
    # We now DEFAULT to generating all formats (safer for Telegram)
    # unless a config file explicitly disables them.
    gen_docx = True
    gen_pdf = True
    
    try:
        config_path = os.path.join('.tmp', 'config.json')
        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                cfg = json.load(f)
                gen_docx = cfg.get("gen_docx", True)
                gen_pdf = cfg.get("gen_pdf", True)
    except Exception as e:
        print(f"Warning: Issue reading config ({e}). Using defaults.")

    # 1. Markdown
    md_path = create_markdown(data, topic, source_url)
    if md_path:
        print(f"Markdown created: {md_path}")
    
    # 2. DOCX
    if gen_docx:
        path = create_docx(data, topic, source_url)
        if path:
            print(f"DOCX created: {path}")
        
    # 3. PDF
    if gen_pdf:
        path = create_pdf(data, topic, source_url)
        if path:
            print(f"PDF created: {path}")


if __name__ == "__main__":
    t = sys.argv[1] if len(sys.argv) > 1 else "Unknown Topic"
    u = sys.argv[2] if len(sys.argv) > 2 else "Unknown Source"
    
    export_reports(t, u)
